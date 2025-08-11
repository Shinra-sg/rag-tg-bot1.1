#!/usr/bin/env python3
import sys
import json
from docx import Document
import os

def parse_word(file_path):
    """Парсит Word файл и возвращает текст и метаданные"""
    try:
        # Загружаем документ
        doc = Document(file_path)
        
        content_parts = []
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            content_parts.append("=== Таблица ===")
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    content_parts.append(row_text)
            content_parts.append("")
        
        content = "\n".join(content_parts)
        
        # Получаем метаданные
        metadata = {
            "title": os.path.basename(file_path),
            "author": None,
            "pages": len(doc.paragraphs) // 20 + 1,  # Примерная оценка страниц
            "created": None,
            "modified": None
        }
        
        # Пытаемся получить метаданные из свойств документа
        if hasattr(doc, 'core_properties'):
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.created:
                metadata["created"] = doc.core_properties.created.isoformat()
            if doc.core_properties.modified:
                metadata["modified"] = doc.core_properties.modified.isoformat()
        
        return {
            "content": content,
            "metadata": metadata
        }
        
    except Exception as e:
        return {
            "content": f"Ошибка парсинга Word файла: {str(e)}",
            "metadata": {"title": os.path.basename(file_path)}
        }

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(json.dumps({"error": "Необходимо указать путь к файлу"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"Файл не найден: {file_path}"}))
        sys.exit(1)
    
    result = parse_word(file_path)
    print(json.dumps(result, ensure_ascii=False, default=str)) 